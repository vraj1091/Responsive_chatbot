# app.py - Enhanced AI Chatbot Backend with File Processing
from ast import Import
import os
import base64
import google.generativeai as genai
from dotenv import load_dotenv
from flask import Flask, json, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_jwt_extended import create_access_token, get_jwt_identity, JWTManager, decode_token
import mysql.connector
from passlib.hash import pbkdf2_sha256 as sha256
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
import mimetypes
import PyPDF2
from PIL import Image
import io
import json
import docx
import traceback

load_dotenv()

# --- CONFIGURATIONS ---
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": ["http://localhost:3000", "http://localhost:5173", "https://responsive-chatbot-2.onrender.com"]}})
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY")
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(hours=1)  # Session management
if not app.config["JWT_SECRET_KEY"]:
    raise RuntimeError("JWT_SECRET_KEY env variable is not set!")
app.config["UPLOAD_FOLDER"] = os.path.join(os.getcwd(), "uploads")
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB max file size

# Create uploads directory if it doesn't exist
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

jwt = JWTManager(app)

# --- Gemini AI Configuration ---
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
print(f"DEBUG: Gemini API key configured: {bool(os.getenv('GEMINI_API_KEY'))}")

# --- Database Configuration ---
db_config = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'user': os.getenv('DB_USER', 'root'),
    'password': os.getenv('DB_PASSWORD', 'your_password'),
    'database': os.getenv('DB_NAME', 'chatbot_db')
}

# Allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'txt', 'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_db_connection():
    return mysql.connector.connect(**db_config)

def validate_token(token):
    """Manually validate JWT token"""
    try:
        decoded_token = decode_token(token)
        return decoded_token['sub']
    except Exception as e:
        print(f"Token validation error: {e}")
        return None

def process_docx_with_gemini(docx_path, user_message):
    """Process DOCX using Gemini AI for question answering"""
    try:
        print(f"DEBUG: Processing DOCX file: {docx_path}")
        
        # Check if file exists
        if not os.path.exists(docx_path):
            print(f"ERROR: File does not exist: {docx_path}")
            return "The uploaded file could not be found."
        
        # Extract text from DOCX
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        extracted_text = '\n'.join(full_text)
        print(f"DEBUG: Extracted text length: {len(extracted_text)}")
        print(f"DEBUG: First 200 chars: {extracted_text[:200]}")
        
        if not extracted_text.strip():
            return "The document appears to be empty or contains no readable text."
        
        # Limit text length for API
        if len(extracted_text) > 8000:
            extracted_text = extracted_text[:8000] + "\n[Document truncated due to length...]"
        
        # Use Gemini to answer based on document content
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""Based on the following document content, please answer the user's question:

DOCUMENT CONTENT:
{extracted_text}

USER QUESTION: {user_message}

Please provide a detailed answer based on the document content. If the information is not available in the document, please state that clearly."""
        
        print(f"DEBUG: Sending to Gemini with prompt length: {len(prompt)}")
        response = model.generate_content(prompt)
        print(f"DEBUG: Received Gemini response: {response.text[:100]}...")
        return response.text
        
    except Exception as e:
        print(f"ERROR processing DOCX: {str(e)}")
        traceback.print_exc()
        return f"Error processing DOCX file: {str(e)}"

def process_txt_with_gemini(txt_path, user_message):
    """Process TXT using Gemini AI for question answering"""
    try:
        print(f"DEBUG: Processing TXT file: {txt_path}")
        
        # Check if file exists
        if not os.path.exists(txt_path):
            print(f"ERROR: File does not exist: {txt_path}")
            return "The uploaded file could not be found."
        
        # Try different encodings
        extracted_text = None
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    extracted_text = f.read()
                print(f"DEBUG: Successfully read file with {encoding} encoding")
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if extracted_text is None:
            return "Could not read the text file due to encoding issues."
        
        print(f"DEBUG: Extracted text length: {len(extracted_text)}")
        print(f"DEBUG: First 200 chars: {extracted_text[:200]}")
        
        if not extracted_text.strip():
            return "The text file appears to be empty."
        
        # Limit text length
        if len(extracted_text) > 8000:
            extracted_text = extracted_text[:8000] + "\n[Document truncated due to length...]"
        
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""Based on the following document content, please answer the user's question:

DOCUMENT CONTENT:
{extracted_text}

USER QUESTION: {user_message}

Please provide a detailed answer based on the document content. If the information is not available in the document, please state that clearly."""
        
        print(f"DEBUG: Sending to Gemini with prompt length: {len(prompt)}")
        response = model.generate_content(prompt)
        print(f"DEBUG: Received Gemini response: {response.text[:100]}...")
        return response.text
        
    except Exception as e:
        print(f"ERROR processing TXT: {str(e)}")
        traceback.print_exc()
        return f"Error processing text file: {str(e)}"

def process_image_with_gemini(image_path, user_message):
    """Process image using Gemini Vision API"""
    try:
        print(f"DEBUG: Processing image file: {image_path}")
        
        if not os.path.exists(image_path):
            return "The uploaded image could not be found."
            
        # Read and encode image
        with open(image_path, 'rb') as image_file:
            image_data = image_file.read()

        # Get mime type
        mime_type, _ = mimetypes.guess_type(image_path)
        if not mime_type or not mime_type.startswith('image/'):
            mime_type = 'image/jpeg'  # Default fallback

        # Use Gemini to analyze image
        model = genai.GenerativeModel('gemini-1.5-flash')

        image_part = {
            "mime_type": mime_type,
            "data": base64.b64encode(image_data).decode()
        }

        prompt = f"User question: {user_message}\n\nPlease analyze this image and provide a detailed response to the user's question."

        response = model.generate_content([prompt, {"inline_data": image_part}])
        return response.text
        
    except Exception as e:
        print(f"ERROR processing image: {str(e)}")
        traceback.print_exc()
        return f"Error processing image: {str(e)}"

def process_pdf_with_gemini(pdf_path, user_message):
    """Process PDF using Gemini"""
    try:
        print(f"DEBUG: Processing PDF file: {pdf_path}")
        
        if not os.path.exists(pdf_path):
            return "The uploaded PDF could not be found."
            
        # Upload PDF to Gemini Files API
        uploaded_file = genai.upload_file(pdf_path)

        # Wait for file to be processed
        import time
        max_wait_time = 60  # Maximum wait time in seconds
        wait_time = 0
        
        while uploaded_file.state.name == "PROCESSING" and wait_time < max_wait_time:
            print(f"Processing PDF... ({wait_time}s)")
            time.sleep(2)
            wait_time += 2
            uploaded_file = genai.get_file(uploaded_file.name)

        if uploaded_file.state.name == "FAILED":
            return "Failed to process PDF file."
            
        if wait_time >= max_wait_time:
            return "PDF processing timed out. Please try with a smaller file."

        # Generate content using the uploaded file
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"User question: {user_message}\n\nPlease analyze this PDF document and provide a detailed response based on its content."

        response = model.generate_content([uploaded_file, prompt])

        # Clean up uploaded file
        try:
            genai.delete_file(uploaded_file.name)
        except:
            pass  # Ignore cleanup errors

        return response.text
        
    except Exception as e:
        print(f"ERROR processing PDF: {str(e)}")
        traceback.print_exc()
        return f"Error processing PDF: {str(e)}"

# --- ROUTES ---
@app.route('/')
def index():
    return jsonify({"message": "AI Chatbot Backend with File Processing", "status": "running"})

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    email = data.get('email', '')

    if not username or not password:
        return jsonify({"msg": "Username and password required"}), 400

    if len(password) < 6:
        return jsonify({"msg": "Password must be at least 6 characters"}), 400

    password_hash = sha256.hash(password)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO users (username, password_hash, email, created_at) VALUES (%s, %s, %s, %s)",
            (username, password_hash, email, datetime.now())
        )
        conn.commit()
    except mysql.connector.IntegrityError:
        return jsonify({"msg": "Username already exists"}), 409
    finally:
        cursor.close()
        conn.close()

    return jsonify({"msg": "User created successfully"}), 201

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    if not username or not password:
        return jsonify({"msg": "Username and password required"}), 400

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

    if user and sha256.verify(password, user['password_hash']):
        access_token = create_access_token(identity=str(user['id']))
        return jsonify({
            "access_token": access_token,
            "user": {
                "id": user['id'],
                "username": user['username'],
                "email": user.get('email', '')
            }
        })

    return jsonify({"msg": "Invalid credentials"}), 401

@app.route('/chat', methods=['POST'])
def chat():
    # Token validation
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({"error": "Authorization header missing"}), 401

    token = auth_header.split(' ')[1]
    current_user_id = validate_token(token)
    if not current_user_id:
        return jsonify({"error": "Invalid token"}), 401

    # Get form data
    user_message = request.form.get('message', '')
    files = request.files.getlist('files')

    print(f"DEBUG: Received message: '{user_message}'")
    print(f"DEBUG: Received {len(files)} files")

    if not user_message and not files:
        return jsonify({"error": "Message or files required"}), 400

    try:
        bot_response = ""
        file_info = []

        # Process uploaded files
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{timestamp}_{filename}"
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                print(f"DEBUG: File saved to: {file_path}")
                print(f"DEBUG: File exists: {os.path.exists(file_path)}")
                if os.path.exists(file_path):
                    print(f"DEBUG: File size: {os.path.getsize(file_path)} bytes")

                file_extension = filename.rsplit('.', 1)[1].lower()
                print(f"DEBUG: Processing file extension: {file_extension}")

                if file_extension == 'pdf':
                    print("DEBUG: Entering PDF processing")
                    if user_message:
                        response = process_pdf_with_gemini(file_path, user_message)
                    else:
                        response = process_pdf_with_gemini(file_path, "Please summarize this document.")
                    bot_response += f"PDF Analysis:\n{response}\n\n"

                elif file_extension in ['png', 'jpg', 'jpeg', 'gif']:
                    print("DEBUG: Entering Image processing")
                    if user_message:
                        response = process_image_with_gemini(file_path, user_message)
                    else:
                        response = process_image_with_gemini(file_path, "Please describe this image.")
                    bot_response += f"Image Analysis:\n{response}\n\n"

                elif file_extension == 'docx':
                    print("DEBUG: Entering DOCX processing")
                    if user_message:
                        response = process_docx_with_gemini(file_path, user_message)
                    else:
                        response = process_docx_with_gemini(file_path, "Please summarize this document.")
                    bot_response += f"DOCX Analysis:\n{response}\n\n"

                elif file_extension == 'txt':
                    print("DEBUG: Entering TXT processing")
                    if user_message:
                        response = process_txt_with_gemini(file_path, user_message)
                    else:
                        response = process_txt_with_gemini(file_path, "Please summarize this document.")
                    bot_response += f"TXT Analysis:\n{response}\n\n"

                file_info.append({
                    "filename": filename,
                    "type": file_extension,
                    "processed": True
                })

                # Clean up file after processing
                try:
                    os.remove(file_path)
                    print(f"DEBUG: File {file_path} cleaned up")
                except Exception as cleanup_error:
                    print(f"DEBUG: Could not cleanup file: {cleanup_error}")

        # If no files, just process text message
        if not files and user_message:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(user_message)
            bot_response = response.text

        if not bot_response:
            bot_response = "I couldn't process your request. Please try again."

        print(f"DEBUG: Final bot response length: {len(bot_response)}")

        # Store chat in database
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO chat_messages (user_id, user_message, bot_response, files_info, created_at) 
            VALUES (%s, %s, %s, %s, %s)
        """, (current_user_id, user_message, bot_response, json.dumps(file_info), datetime.now()))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({
            "reply": bot_response,
            "files_processed": len(file_info)
        })

    except Exception as e:
        print(f"ERROR in chat endpoint: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to process request: {str(e)}"}), 500

@app.route('/history', methods=['GET'])
def get_history():
    # Token validation
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({"error": "Authorization header missing"}), 401

    token = auth_header.split(' ')[1]
    current_user_id = validate_token(token)
    if not current_user_id:
        return jsonify({"error": "Invalid token"}), 401

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT id, user_message, bot_response, files_info, created_at 
            FROM chat_messages 
            WHERE user_id = %s 
            ORDER BY created_at DESC 
            LIMIT 50
        """, (current_user_id,))

        messages = cursor.fetchall()
        cursor.close()
        conn.close()

        # Parse JSON fields and format dates
        for message in messages:
            message['files_info'] = json.loads(message['files_info'] or '[]')
            message['created_at'] = message['created_at'].isoformat()

        return jsonify({"messages": messages})

    except Exception as e:
        print(f"Error getting history: {e}")
        return jsonify({"error": "Failed to get history"}), 500

@app.route('/clear-history', methods=['DELETE'])
def clear_history():
    # Token validation
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({"error": "Authorization header missing"}), 401

    token = auth_header.split(' ')[1]
    current_user_id = validate_token(token)
    if not current_user_id:
        return jsonify({"error": "Invalid token"}), 401

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM chat_messages WHERE user_id = %s", (current_user_id,))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({"msg": "History cleared successfully"})

    except Exception as e:
        print(f"Error clearing history: {e}")
        return jsonify({"error": "Failed to clear history"}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  # For Render deployment
    app.run(host="0.0.0.0", port=port, debug=True)
